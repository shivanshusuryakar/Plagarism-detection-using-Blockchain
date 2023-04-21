from give_hash import generate_hash
from docx import Document
from PyPDF2 import PdfReader


def give_plag_content(filename,plag_list):
	final_hash=[]
	if filename.split(".")[1]=="pdf":
		p=PdfReader(filename)
		for i in range(len(p.pages)):
			page=p.pages[i]
			text=page.extract_text()
			text=text.strip()
			text=text.split(".")
			for x in text:
				final_hash.append(x)

	if filename.split(".")[1]=="docx" or filename.split(".")[1]=="doc":
		d=Document(filename)
		for i in range(len(d.paragraphs)):
			x=d.paragraphs[i].text
			x=x.strip()
			final_hash.append(x)
	else:

		with open(filename,"r",encoding="utf-8") as f:
			x=f.read()
			x=x.strip()
			x=x.replace('\n','')
			l=x.split('.')
			for x in l:
				x=x.strip()
				if x=='':
					continue
				final_hash.append(x)

	new_plag_content=[]
	for i in range(len(final_hash)):
		if i in plag_list:
			x="<span style='color:#c44c37'>"+final_hash[i]+"."+"</span>"
		else:
			x="<span style='color:gray'>"+final_hash[i]+"."+"</span>"
		new_plag_content.append(x)

	final_doc=""
	for x in new_plag_content:
		final_doc+=x
		
	return final_doc

