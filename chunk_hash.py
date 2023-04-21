from give_hash import generate_hash
from docx import Document
from PyPDF2 import PdfReader

def give_chunk_hash(filename):
	final_hash=[]
	if filename.split(".")[1]=="pdf":
		p=PdfReader(filename)
		for i in range(len(p.pages)):
			page=p.pages[i]
			text=page.extract_text()
			text=text.strip()
			text=text.split(".")
			for x in text:
				final_hash.append(generate_hash(x))
		return final_hash

	if filename.split(".")[1]=="docx" or filename.split(".")[1]=="doc":
		d=Document(filename)
		for i in range(len(d.paragraphs)):
			x=d.paragraphs[i].text
			x=x.strip()
			final_hash.append(generate_hash(x))
		return final_hash
	else:

		with open(filename,"r",encoding="utf-8") as f:
			x=f.read()
			x=x.strip()
			x=x.replace('\n','')
			l=x.split('.')
			for x in l:
				x=x.strip()
				if x=='':
					continue
				final_hash.append(generate_hash(x))
			return final_hash

